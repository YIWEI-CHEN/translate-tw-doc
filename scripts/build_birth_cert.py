# -*- coding: utf-8 -*-
"""
Taiwan Birth Certificate (出生證明書) — English Translation Formatter
=====================================================================
Generates a formatted Word (.docx) translation of a Taiwan ROC birth
certificate, including a translator certification page and appended
original-document scans.

Usage:
    uv run --link-mode=copy --with python-docx --with pymupdf \
        build_birth_cert.py

Requires: python-docx, pymupdf (for rendering PDF → PNG if needed)

License: MIT
"""
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Symbol helpers ──
B = chr(9632)   # ■ filled box
E = chr(9633)   # □ empty box
C = lambda n: chr(0x245F + n)   # circled number ①②③...

# ─────────────────────────────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────────────────────────────

def set_table_borders(table, sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    ex = tblPr.find(qn('w:tblBorders'))
    if ex is not None:
        tblPr.remove(ex)
    tblBorders = OxmlElement('w:tblBorders')
    for name in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def ct(cell, text, bold=False, size=9, align='left'):
    amap = {'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT}
    para = cell.paragraphs[0]
    para.alignment = amap.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

def ct2(cell, line1, line2, bold1=True, size1=13, size2=9, align='center'):
    """Two lines in one cell."""
    amap = {'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER}
    para = cell.paragraphs[0]
    para.alignment = amap.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    para.clear()
    r1 = para.add_run(line1)
    r1.bold = bold1
    r1.font.size = Pt(size1)
    r1.add_break()
    r2 = para.add_run(line2)
    r2.bold = False
    r2.font.size = Pt(size2)

# ─────────────────────────────────────────────────────────────────────
# FIELD FORMATTERS
# ─────────────────────────────────────────────────────────────────────

def fmt_birth_type(type_num, male_count, female_count, birth_order=None,
                   female_selected=None, female_display=None):
    """Format (8) Birth Type with filled/empty checkbox boxes.

    Args:
        type_num:         1=Single, 2=Twins, 3=Triplets, 4=Other
        male_count:       number of males in this delivery
        female_count:     number of females in this delivery
        birth_order:      birth order if multiple birth (or None)
        female_selected:  override checkbox state (True/False);
                          defaults to female_count > 0
        female_display:   override display text (e.g. '__');
                          defaults to str(female_count)
    """
    types = [(1,"Single birth"),(2,"Twins"),(3,"Triplets"),(4,"Other multiple births")]
    top = "  ".join(f"{C(n)} {label} {B if n==type_num else E}" for n,label in types)
    m_box = B if male_count > 0 else E
    m_disp = str(male_count) if male_count > 0 else "__"
    f_box = (B if female_selected else E) if female_selected is not None else (B if female_count > 0 else E)
    f_disp = female_display if female_display is not None else (str(female_count) if female_count > 0 else "__")
    order_box = B if (birth_order and birth_order != '___') else E
    mid = (f"  Male {m_box} {m_disp} person(s)   "
           f"Female {f_box} {f_disp} person(s)   "
           f"Birth Order in This Delivery {order_box}")
    return top + "\n" + mid

def fmt_birth_location(type_num, address):
    """Format (10) Birth Location with filled/empty boxes + address.

    type_num: 1=Hospital, 2=Clinic, 3=Birthing Center, 4=Home, 5=Other
    """
    types = [(1,"Hospital"),(2,"Clinic"),(3,"Birthing Center"),(4,"Home"),(5,"Other")]
    boxes = "  ".join(f"{C(n)} {label} {B if n==type_num else E}" for n,label in types)
    return boxes + "\n  Address:  " + address

def fmt_delivery(type_num, seal_note=None):
    """Format (11) Delivery Attended By with checkboxes.

    type_num: 1=Physician, 2=Midwife, 3=Other
    seal_note: optional annotation e.g. '[Illegible seal]'
    """
    types = [(1,"Physician"),(2,"Midwife"),(3,"Other")]
    parts = []
    for n, label in types:
        part = f"{C(n)} {label} {B if n==type_num else E}"
        if seal_note and n == type_num:
            part += f"  {seal_note}"
        parts.append(part)
    return "  ".join(parts)

# ─────────────────────────────────────────────────────────────────────
# SHARED: TRANSLATOR CERT + ORIGINAL PAGES
# ─────────────────────────────────────────────────────────────────────

def add_translator_cert(doc, doc_title):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Certification by Translator')
    r.bold = True; r.font.size = Pt(13)
    doc.add_paragraph()
    doc.add_paragraph().add_run(
        "I, [NAME OF TRANSLATOR] ________________________________________,"
        " certify that I am fluent in the English and Chinese (Traditional) "
        "languages, and that the above document is an accurate translation of the document attached "
        "entitled: " + doc_title + "."
    ).font.size = Pt(11)
    doc.add_paragraph()
    for lbl in ["Date:", "Signature:", "", "", "[Typed Full Name]", "", "[Address]"]:
        doc.add_paragraph().add_run(lbl).font.size = Pt(11)

def add_original_pages(doc, img_paths):
    for i, img_path in enumerate(img_paths):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = "ORIGINAL DOCUMENT - Page " + str(i+1)
        if i == 0:   label += " (Front)"
        elif i == 1: label += " (Back / Certification)"
        r = p.add_run(label); r.bold = True; r.font.size = Pt(11)
        doc.add_paragraph()
        pi = doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.add_run().add_picture(img_path, width=Inches(5.5))

# ─────────────────────────────────────────────────────────────────────
# BIRTH CERTIFICATE BUILDER
# Columns: (1)Relationship | Name | DOB | Place of Origin | Household Reg.
# ─────────────────────────────────────────────────────────────────────
def build_birth_cert(d):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin   = Inches(0.9)
    sec.right_margin  = Inches(0.9)

    # 5 cols: 0.55 + 1.45 + 1.15 + 1.2 + 2.35 = 6.7"
    W = [Inches(0.55), Inches(1.45), Inches(1.15), Inches(1.2), Inches(2.35)]
    NC = 5

    table = doc.add_table(rows=0, cols=NC)
    set_table_borders(table)

    def sw(row):
        for i,w in enumerate(W): row.cells[i].width = w

    def full_row(text, bold=False, size=9, align='left', height=None, bg=None):
        row = table.add_row()
        if height: row.height = Cm(height)
        sw(row)
        cell = row.cells[0].merge(row.cells[NC-1])
        ct(cell, text, bold=bold, size=size, align=align)
        if bg: set_cell_bg(cell, bg)
        return cell

    # ── TITLE ROW: full width ──
    row = table.add_row()
    row.height = Cm(1.1)
    sw(row)
    title_cell = row.cells[0].merge(row.cells[4])
    ct2(title_cell,
        "BIRTH CERTIFICATE",
        "(English Translation of 出生證明書)",
        bold1=True, size1=13, size2=9, align='center')
    set_cell_bg(title_cell, 'D9E1F2')

    # ── BIRTH CERTIFICATE NO. ROW ──
    row = table.add_row()
    row.height = Cm(0.55)
    sw(row)
    cert_cell = row.cells[0].merge(row.cells[4])
    ct(cert_cell, "Birth Certificate No.: ___________", size=9, align='right')

    # ── PARENTS HEADER ──
    row = table.add_row()
    row.height = Cm(0.65)
    sw(row)
    ct(row.cells[0], "(1) Relationship",              bold=True, size=7.5)
    ct(row.cells[1], "Name",                          bold=True, size=7.5, align='center')
    ct(row.cells[2], "Date of Birth",                 bold=True, size=7.5, align='center')
    ct(row.cells[3], "Place of Origin",               bold=True, size=7.5)
    ct(row.cells[4], "Household Registration Address",bold=True, size=7.5)
    for c in row.cells: set_cell_bg(c, 'F2F2F2')

    # ── FATHER ──
    row = table.add_row()
    sw(row)
    ct(row.cells[0], "Father", bold=True, size=9)
    ct(row.cells[1], d['father_name'],                   size=9)
    ct(row.cells[2], d['father_dob'],                    size=8)
    ct(row.cells[3], d.get('father_place_of_origin',''), size=8)
    ct(row.cells[4], d['household_reg'],                 size=8)

    # ── MOTHER ──
    row = table.add_row()
    sw(row)
    ct(row.cells[0], "Mother", bold=True, size=9)
    ct(row.cells[1], d['mother_name'],                   size=9)
    ct(row.cells[2], d['mother_dob'],                    size=8)
    ct(row.cells[3], d.get('mother_place_of_origin',''), size=8)
    ct(row.cells[4], d['household_reg'],                 size=8)

    # ── (2) IDs ──
    full_row(
        "(2)  Father's & Mother's National ID No.     "
        "Father:  " + d['father_id'] + "          Mother:  " + d['mother_id'],
        size=9)

    # ── (3) OCCUPATION ──
    full_row(
        "(3)  Parents' Occupation\n"
        "    " + C(1) + " Father's workplace & occupation:  " + d['father_occ'] + "\n"
        "    " + C(2) + " Mother's workplace & occupation:  " + d['mother_occ'],
        size=9)

    # ── (4)(5)(6)(7) HEADER ──
    row = table.add_row()
    row.height = Cm(0.9)
    sw(row)
    sx = row.cells[0].merge(row.cells[1])
    ct(sx, "(4) Sex of the Newborn", bold=True, size=8)
    ct(row.cells[2],
       "(5) Total Live Births\nby Mother\n(including this\nlive birth)",
       bold=True, size=7.5)
    ct(row.cells[3], "(6) Gestational\nAge (Weeks)", bold=True, size=8)
    ct(row.cells[4],
       "(7) Birth Weight\n(leave blank if\nnot weighed)",
       bold=True, size=8)
    for c in row.cells: set_cell_bg(c, 'F2F2F2')

    # ── (4)(5)(6)(7) VALUES ──
    row = table.add_row()
    sw(row)
    sx = row.cells[0].merge(row.cells[1])
    m_box  = B if d['sex'] == 'Male'   else E
    f_box  = B if d['sex'] == 'Female' else E
    ct(sx, C(1) + " Male "   + m_box + "   " +
           C(2) + " Female " + f_box, size=9)
    ct(row.cells[2], str(d.get('total_live_births','N/A')), size=9)
    ct(row.cells[3], d['gestational_age'],                  size=9)
    ct(row.cells[4], d['birth_weight'],                     size=8)

    # ── (8) BIRTH TYPE ──
    full_row(
        "(8)  Birth Type:\n  " +
        fmt_birth_type(d['birth_type_num'],
                       d['birth_male_count'],
                       d['birth_female_count'],
                       d.get('birth_order_in_multiple'),
                       d.get('birth_female_selected'),
                       d.get('birth_female_display')),
        size=9)

    # ── (9) DATE AND TIME ──
    full_row("(9)  Date and Time of Birth:   " + d['dob_time'], size=9)

    # ── (10) BIRTH LOCATION + ADDRESS ──
    full_row(
        "(10)  Birth Location:\n  " +
        fmt_birth_location(d['birth_location_type_num'],
                           d['birth_location_addr']),
        size=9)

    # ── (11) DELIVERY ──
    full_row(
        "(11)  Delivery Attended By:   " +
        fmt_delivery(d['delivery_by_num'], d.get('delivery_seal_note')),
        size=9)

    # ── CHILD ID + CHILD NAME ──
    if d.get('child_name'):
        row = table.add_row()
        sw(row)
        left = row.cells[0].merge(row.cells[2])
        ct(left, "Child's National ID No.:   " + d['child_id'], size=9, align='left')
        right = row.cells[3].merge(row.cells[4])
        ct(right, "Child's Name:   " + d['child_name'], size=9, align='left')
    else:
        full_row("Child's National ID No.:   " + d['child_id'], size=9)

    # ── (12) SPECIAL CONDITIONS ──
    full_row(
        "(12)  Special Medical Conditions of the mother before/after delivery\n"
        "      and of the infant before/after birth:   " + d['special_conditions'],
        size=9)

    # ── CERTIFICATION STATEMENT ──
    full_row(
        "The above information is true and correct. This certificate is hereby issued.",
        bold=True, size=9, bg='FFF2CC')

    # ── PHYSICIAN / MIDWIFE NAME ROW ──
    row = table.add_row()
    sw(row)
    lbl = row.cells[0].merge(row.cells[1])
    ct(lbl, "Physician's or Midwife's Name:", bold=True, size=8)
    val = row.cells[2].merge(row.cells[4])
    ct(val, d['physician_name'], size=9)

    # ── HOSPITAL INFO ──
    for label, key in [
        ("Physician or Midwife\nCertificate No.:", 'med_cert_no'),
        ("Hospital, Clinic, or\nBirthing Center Name:",  'hospital_name'),
        ("Medical Facility\nLicense No.:",         'med_license_no'),
        ("Medical Facility\nAddress:",             'med_facility_addr'),
    ]:
        row = table.add_row()
        sw(row)
        lbl = row.cells[0].merge(row.cells[1])
        ct(lbl, label, bold=True, size=8)
        val = row.cells[2].merge(row.cells[4])
        ct(val, d[key], size=9)

    row = table.add_row()
    sw(row)
    lbl = row.cells[0].merge(row.cells[1])
    ct(lbl, "Physician's Signature Date:", bold=True, size=8)
    val = row.cells[2].merge(row.cells[4])
    ct(val, d['physician_date'], size=9)

    # ── NOTICE ──
    full_row(
        "Household registration must be filed within 15 days of birth to avoid penalty."
        "                                                        (19 x 27 cm)",
        size=8, bg='F2F2F2')

    # ── PAGE 2 — CERTIFICATION OF COPY ──
    full_row("PAGE 2 — CERTIFICATION OF COPY",
             bold=True, size=10, align='center', bg='D9E1F2')
    for label, key in [
        ("Certification Statement", 'p2_cert_statement'),
        ("Issuing Office",          'p2_office'),
        ("Director",                'p2_director'),
        ("Date of Certification",   'p2_date'),
        ("Document Serial No.",     'p2_serial'),
    ]:
        row = table.add_row()
        sw(row)
        lbl = row.cells[0].merge(row.cells[1])
        ct(lbl, label, bold=True, size=9)
        set_cell_bg(lbl, 'F2F2F2')
        val = row.cells[2].merge(row.cells[4])
        ct(val, d[key], size=9)

    return doc

# ─────────────────────────────────────────────────────────────────────
# EXAMPLE DATA  (replace with your own translated values)
# ─────────────────────────────────────────────────────────────────────
# Name format:  SURNAME, GIVEN-NAME  (ALL CAPS, Wade-Giles romanization)
# Date format:  Month Day, Year (ROC Year ##)
#               ROC Year = Western Year - 1911
# Address:      translated to English, from smallest to largest unit

example_birth = dict(
    father_name            = "WANG, DA-MING",
    father_dob             = "January 15, 1960 (ROC Year 49)",
    father_id              = "A123456789",
    father_occ             = "N/A",
    father_place_of_origin = "Taipei City, Taiwan Province",
    mother_name            = "LIN, MEI-HUA",
    mother_dob             = "March 20, 1962 (ROC Year 51)",
    mother_id              = "B223456789",
    mother_occ             = "N/A",
    mother_place_of_origin = "Kaohsiung City, Taiwan Province",
    household_reg          = "No. 100, Section 2, Zhongshan Road, Zhongzheng District, Taipei City",
    child_id               = "A234567890",
    # child_name           = "WANG, XIAO-MING  [Illegible seal]",   # optional: splits cell
    sex                    = "Male",
    dob_time               = "June 1, 1990 (ROC Year 79), at 10:30 AM",
    total_live_births      = "1",
    gestational_age        = "39 weeks (full term)",
    birth_weight           = "3,200 g",
    birth_type_num         = 1,      # 1=Single, 2=Twins, 3=Triplets, 4=Other
    birth_male_count       = 1,
    birth_female_count     = 0,
    birth_order_in_multiple= None,
    birth_location_type_num= 1,      # 1=Hospital, 2=Clinic, 3=Birthing Center, 4=Home, 5=Other
    birth_location_addr    = "No. 7, Zhongshan South Road, Zhongzheng District, Taipei City",
    delivery_by_num        = 1,      # 1=Physician, 2=Midwife, 3=Other
    # delivery_seal_note   = "[Illegible seal]",   # optional: annotation next to checkbox
    special_conditions     = "Normal",
    physician_name         = "CHEN, YI-SHENG",
    hospital_name          = "Example Hospital",
    med_cert_no            = "Example-Health-Yi-001",
    med_license_no         = "Example-Health-Zhi-001",
    med_facility_addr      = "No. 7, Zhongshan South Road, Zhongzheng District, Taipei City",
    physician_date         = "June 3, 1990 (ROC Year 79)",
    p2_cert_statement      = "This photocopy corresponds to the original document.",
    p2_office              = "Example Household Registration Office",
    p2_director            = "ZHANG, ZHU-REN  [Official seal]",
    p2_date                = "January 1, 2026 (ROC Year 115)",
    p2_serial              = "Example District Household Transcript No. (A) 000001",
)

# ─────────────────────────────────────────────────────────────────────
# BUILD & SAVE
# ─────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    import os

    # Set your output directory here
    BASE = os.path.dirname(os.path.abspath(__file__))

    doc = build_birth_cert(example_birth)
    add_translator_cert(doc, "Birth Certificate (出生證明書)")
    # Uncomment to append original document scans:
    # add_original_pages(doc, [
    #     os.path.join(BASE, "birth_cert_page1.png"),
    #     os.path.join(BASE, "birth_cert_page2.png"),
    # ])
    out = os.path.join(BASE, "Birth Certificate - English Translation (Formatted).docx")
    doc.save(out)
    print("Saved:", os.path.basename(out))
