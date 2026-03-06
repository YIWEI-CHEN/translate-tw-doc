# -*- coding: utf-8 -*-
"""
Taiwan Military Discharge Certificate (替代役退役證明書) — English Translation Formatter
======================================================================================
Generates a formatted Word (.docx) translation of a Taiwan ROC substitute
military service discharge certificate (or replacement certificate),
including a translator certification page and appended original-document scans.

Usage:
    uv run --link-mode=copy --with python-docx --with pymupdf \
        build_military.py

Requires: python-docx, pymupdf (for rendering PDF → PNG if needed)

License: MIT
"""
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────────────────────────────

def set_table_borders(table, sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    ex = tblPr.find(qn('w:tblBorders'))
    if ex is not None:
        tblPr.remove(ex)
    tblBorders = OxmlElement('w:tblBorders')
    for name in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def ct(cell, text, bold=False, size=9, align='left'):
    amap = {'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT}
    para = cell.paragraphs[0]
    para.alignment = amap.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

# ─────────────────────────────────────────────────────────────────────
# SHARED: TRANSLATOR CERT + ORIGINAL PAGES
# ─────────────────────────────────────────────────────────────────────

def add_translator_cert(doc, doc_title):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Certification by Translator')
    r.bold = True; r.font.size = Pt(13)
    doc.add_paragraph()
    doc.add_paragraph().add_run(
        "I, [NAME OF TRANSLATOR] ________________________________________,"
        " certify that I am fluent in the English and Chinese (Traditional) "
        "languages, and that the above document is an accurate translation of the document attached "
        "entitled: " + doc_title + "."
    ).font.size = Pt(11)
    doc.add_paragraph()
    for lbl in ["Date:", "Signature:", "", "", "[Typed Full Name]", "", "[Address]"]:
        doc.add_paragraph().add_run(lbl).font.size = Pt(11)

def add_original_pages(doc, img_paths):
    for i, img_path in enumerate(img_paths):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = "ORIGINAL DOCUMENT - Page " + str(i+1)
        r = p.add_run(label); r.bold = True; r.font.size = Pt(11)
        doc.add_paragraph()
        pi = doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pi.add_run().add_picture(img_path, width=Inches(5.5))

# ─────────────────────────────────────────────────────────────────────
# MILITARY RECORD BUILDER
# ─────────────────────────────────────────────────────────────────────
def build_military(d):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.5)
    sec.right_margin  = Inches(1.5)

    W = [Inches(2.3), Inches(3.2)]
    table = doc.add_table(rows=0, cols=2)
    set_table_borders(table)

    def sw(row):
        row.cells[0].width = W[0]; row.cells[1].width = W[1]

    def full_row(text, bold=False, size=10, align='center', height=None, bg=None):
        row = table.add_row()
        if height: row.height = Cm(height)
        sw(row)
        cell = row.cells[0].merge(row.cells[1])
        ct(cell, text, bold=bold, size=size, align=align)
        if bg: set_cell_bg(cell, bg)

    def data_row(label, value, size=10):
        row = table.add_row()
        sw(row)
        ct(row.cells[0], label, bold=True, size=size)
        set_cell_bg(row.cells[0], 'F2F2F2')
        ct(row.cells[1], value, size=size)

    full_row("REPLACEMENT CERTIFICATE FOR\nSUBSTITUTE MILITARY SERVICE DISCHARGE",
             bold=True, size=11, height=1.2, bg='D9E1F2')
    full_row("(English Translation of 替代役退役證明書補發證明)", size=8, height=0.55)

    data_row("Name",                      d['name'])
    data_row("Date of Birth",             d['dob'])
    data_row("National ID No.",           d['id_no'])
    data_row("Enlistment Date",           d['enlistment'])
    data_row("Discharge Effective Date",  d['discharge'])
    data_row("Service Type",              d['service_type'])
    data_row("Supervisory Cadre",         d['service_unit'])
    data_row("Discharge Certificate No.", d['discharge_cert'])
    data_row("Reason for Discharge",      d['reason'])

    row = table.add_row()
    row.height = Cm(2.0)
    sw(row)
    ct(row.cells[0], "Authorizing Official\n(Official Seal)", bold=True, size=10)
    set_cell_bg(row.cells[0], 'F2F2F2')
    ct(row.cells[1], "[Official Seal & Signature]\n\nDistrict Director: " + d['official'], size=10)

    full_row("Republic of China (Taiwan)     " + d['issue_date'] + "\n"
             "Document Reference No.: " + d['ref_no'],
             size=9, height=1.0)

    return doc

# ─────────────────────────────────────────────────────────────────────
# EXAMPLE DATA  (replace with your own translated values)
# ─────────────────────────────────────────────────────────────────────
# Name format:  SURNAME, GIVEN-NAME  (ALL CAPS, Wade-Giles romanization)
# Date format:  Month Day, Year (ROC Year ##)
#               ROC Year = Western Year - 1911

example_military = dict(
    name           = "WANG, DA-MING",
    dob            = "June 1, 1990 (ROC Year 79)",
    id_no          = "A123456789",
    enlistment     = "July 1, 2013 (ROC Year 102)",
    discharge      = "Effective from midnight, June 30, 2014 (ROC Year 103)",
    service_type   = "Educational Service Corps",
    service_unit   = "Non-supervisory Cadre",
    discharge_cert = "No. 10300001",
    reason         = "Completion of Service Term",
    official       = "ZHANG, ZHU-REN",
    issue_date     = "January 1, 2026 (ROC Year 115)",
    ref_no         = "Example-District-Yi No. 001 (Military Service)",
)

# ─────────────────────────────────────────────────────────────────────
# BUILD & SAVE
# ─────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    import os

    # Set your output directory here
    BASE = os.path.dirname(os.path.abspath(__file__))

    doc = build_military(example_military)
    add_translator_cert(doc, "Replacement Certificate for Substitute Military Service Discharge")
    # Uncomment to append original document scan:
    # add_original_pages(doc, [os.path.join(BASE, "military_page1.png")])
    out = os.path.join(BASE, "Military Record - English Translation (Formatted).docx")
    doc.save(out)
    print("Saved:", os.path.basename(out))
